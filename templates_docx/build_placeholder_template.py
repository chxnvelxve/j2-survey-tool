"""Build templates_docx/survey_report.docx — run to regenerate the placeholder template.

Phase 9 shell: all 8 canonical sections. Bind against the frozen context contract
in docs/template_map.md. 🔒 Layout assumptions until Josh's sample deliverable.

Jinja tag inventory:
  company_name, primary_color, logo, job_name, project_name
  survey_type, location_vertical, band_plan, site_metadata
  success_criteria.{profile_key,label,min_rssi_dbm,min_snr_db,
                    min_data_rate_mbps,max_co_channel_aps,is_override}
  exec_summary, scope_methodology, findings, ap_count, override_count
  aps[]: name, model, vendor, floor, x, y, status, radios_summary,
         photo_close, photo_far, photo_close_label, photo_far_label
  attachments[]: filename, image, is_image
  overrides[]: ap_name, type_label, detail, reason
  has_overrides
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).with_name("survey_report.docx")


def _heading(doc: Document, text: str, *, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def build() -> None:
    doc = Document()

    # 1. Cover — AUTO
    _para(doc, "{{ logo }}")
    _heading(doc, "{{ company_name }}", level=0)
    _para(doc, "Job: {{ job_name }}")
    _para(doc, "Project: {{ project_name }}")
    _para(doc, "Brand color: {{ primary_color }}")

    # 2. Executive Summary — DRAFTED
    _heading(doc, "Executive Summary")
    _para(doc, "{{ exec_summary }}")

    # 3. Scope / Methodology — DRAFTED + AUTO crumbs
    _heading(doc, "Scope / Methodology")
    _para(doc, "{{ scope_methodology }}")
    _para(doc, "Survey type: {{ survey_type }}")
    _para(doc, "Location vertical: {{ location_vertical }}")
    _para(doc, "Band plan: {{ band_plan }}")
    _para(doc, "Site metadata: {{ site_metadata }}")

    # 4. Success Criteria — AUTO
    _heading(doc, "Success Criteria")
    _para(doc, "Profile: {{ success_criteria.label }} ({{ success_criteria.profile_key }})")
    _para(doc, "Min RSSI: {{ success_criteria.min_rssi_dbm }} dBm")
    _para(doc, "Min SNR: {{ success_criteria.min_snr_db }} dB")
    _para(doc, "Min data rate: {{ success_criteria.min_data_rate_mbps }} Mbps")
    _para(doc, "Max co-channel overlap: {{ success_criteria.max_co_channel_aps }} APs")
    _para(doc, "Override applied: {{ success_criteria.is_override }}")

    # 5. Findings — DRAFTED + AUTO counts
    _heading(doc, "Findings")
    _para(doc, "{{ findings }}")
    _para(doc, "Access points in inventory: {{ ap_count }}")
    _para(doc, "Overridden issues: {{ override_count }}")

    # 6. AP Inventory — AUTO (photo pair layout)
    _heading(doc, "AP Inventory")
    _para(doc, "{%p for ap in aps %}")
    _para(doc, "{{ ap.name }} — {{ ap.model }} — {{ ap.floor }} — {{ ap.status }}")
    _para(doc, "Vendor: {{ ap.vendor }}  |  Position: X={{ ap.x }} Y={{ ap.y }}")
    _para(doc, "Radios: {{ ap.radios_summary }}")
    _para(doc, "Close: {{ ap.photo_close_label }}")
    _para(doc, "{{ ap.photo_close }}")
    _para(doc, "Far: {{ ap.photo_far_label }}")
    _para(doc, "{{ ap.photo_far }}")
    _para(doc, "{%p endfor %}")

    # 7. Issues & Gaps — AUTO
    _heading(doc, "Issues & Gaps")
    _para(doc, "{% if has_overrides %}")
    _para(doc, "{%p for ov in overrides %}")
    _para(doc, "{{ ov.ap_name }} — {{ ov.type_label }}: {{ ov.detail }} ({{ ov.reason }})")
    _para(doc, "{%p endfor %}")
    _para(doc, "{% else %}")
    _para(doc, "No overridden issues recorded for this job.")
    _para(doc, "{% endif %}")

    # 8. Appendices — AUTO
    _heading(doc, "Appendices")
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Filename"
    hdr[1].text = "Preview"
    _para(doc, "{%p for att in attachments %}")
    _para(
        doc,
        "{% if att.image %}{{ att.filename }} | {{ att.image }}"
        "{% else %}{{ att.filename }} | (file){% endif %}",
    )
    _para(doc, "{%p endfor %}")

    run = doc.add_paragraph().add_run("Prepared by {{ company_name }}")
    run.font.size = Pt(9)

    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
