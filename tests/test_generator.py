"""Unit tests for the generator stage."""
from __future__ import annotations

import logging
from pathlib import Path

import pytest
from docx import Document
from docxtpl import DocxTemplate, InlineImage

from app.models.enums import PhotoShotType
from app.services.generator.context import build_template_context
from app.schemas.merge import (
    Flag,
    FlagType,
    MergedAP,
    MergedAPStatus,
    MergedJob,
    MergedPhotoRef,
    MergedPhotoSlots,
)
from app.schemas.survey import SurveyAP, SurveyRadio
from app.services.generator.errors import (
    EmptyMergedJobError,
    MissingTemplateError,
    NoAttachmentsError,
    UnresolvedFlagsError,
)
from app.services.generator.generator import generate_docx
from app.services.generator.types import AttachmentInput, BrandingConfig

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = ROOT / "templates_docx" / "survey_report.docx"
LOGO = ROOT / "branding" / "j2_logo_placeholder.png"
FIXTURES = Path(__file__).parent / "fixtures"
PHOTO_CLOSE = FIXTURES / "gen_photo_close.png"
PHOTO_FAR = FIXTURES / "gen_photo_far.png"


def _ensure_photos() -> None:
    if not PHOTO_CLOSE.exists():
        from tests.fixtures.build_gen_photos import MINIMAL_PNG

        PHOTO_CLOSE.write_bytes(MINIMAL_PNG)
        PHOTO_FAR.write_bytes(MINIMAL_PNG)


def _merged_job(*, with_override: bool = False) -> MergedJob:
    ap = MergedAP(
        ap_name="AP-01-NE",
        survey_data=SurveyAP(
            name="AP-01-NE",
            model="C9136I",
            vendor="Cisco",
            floor_id="fp1",
            x=12.5,
            y=34.0,
            radios=[SurveyRadio(band="5GHz", channel=36, tx_power=17.0)],
        ),
        photos=MergedPhotoSlots(
            close=MergedPhotoRef(photo_id=1, original_filename="close.png"),
            far=MergedPhotoRef(photo_id=2, original_filename="far.png"),
        ),
        status=MergedAPStatus.MATCHED,
    )
    flags: list[Flag] = []
    if with_override:
        flags.append(
            Flag(
                ap_name="AP-02-SW",
                type=FlagType.MISSING_PHOTO,
                detail="Missing far photo",
                override_reason="Ceiling access denied",
            ),
        )
    return MergedJob(aps=[ap], flags=flags)


def _attachments(tmp_path: Path) -> list[AttachmentInput]:
    from tests.fixtures.build_gen_photos import MINIMAL_PNG

    att_path = tmp_path / "idf.png"
    att_path.write_bytes(MINIMAL_PNG)
    return [
        AttachmentInput(
            attachment_id=1,
            filename="idf.png",
            path=att_path,
            content_type="image/png",
        ),
    ]


SECTION_HEADINGS = (
    "Executive Summary",
    "Scope / Methodology",
    "Success Criteria",
    "Findings",
    "AP Inventory",
    "Issues & Gaps",
    "Appendices",
)


def _branding() -> BrandingConfig:
    return BrandingConfig(
        company_name="Test Survey Co",
        logo_path=None,
        primary_color="#1F4E79",
    )


def test_generate_docx_happy_path(tmp_path: Path) -> None:
    _ensure_photos()
    merged = _merged_job(with_override=True)
    docx_bytes = generate_docx(
        merged,
        template_path=TEMPLATE,
        branding=_branding(),
        job_name="Site Alpha",
        project_name="Test Building Survey",
        floor_name_for=lambda _fid: "Floor 1",
        photo_paths={1: PHOTO_CLOSE, 2: PHOTO_FAR},
        attachments=_attachments(tmp_path),
        location_vertical="warehouse",
        survey_type="validation",
        band_plan="5GHz primary",
        site_metadata="Building A",
    )

    out = tmp_path / "out.docx"
    out.write_bytes(docx_bytes)
    doc = Document(out)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Site Alpha" in text
    assert "Test Building Survey" in text
    assert "AP-01-NE" in text
    assert "Ceiling access denied" in text
    assert "Test Survey Co" in text
    for heading in SECTION_HEADINGS:
        assert heading in text, f"Missing section heading: {heading}"
    assert "Warehouse" in text
    assert "-67" in text
    assert "144" in text
    assert "validation" in text
    assert "Building A" in text


def test_logo_inline_image_present_when_configured(tmp_path: Path) -> None:
    _ensure_photos()
    assert LOGO.is_file(), "Canonical placeholder logo missing — run build_placeholder_logo.py"
    tpl = DocxTemplate(str(TEMPLATE))
    ctx = build_template_context(
        tpl,
        _merged_job(),
        branding=BrandingConfig(
            company_name="Test Survey Co",
            logo_path=LOGO,
            primary_color="#1F4E79",
        ),
        job_name="Job",
        project_name="Project",
        floor_name_for=lambda _fid: "Floor 1",
        photo_paths={1: PHOTO_CLOSE, 2: PHOTO_FAR},
        attachments=_attachments(tmp_path),
    )
    assert isinstance(ctx["logo"], InlineImage)


def test_unresolvable_logo_path_warns_and_omits(
    tmp_path: Path,
    caplog: pytest.LogCaptureFixture,
) -> None:
    _ensure_photos()
    tpl = DocxTemplate(str(TEMPLATE))
    with caplog.at_level(logging.WARNING, logger="app.services.generator.context"):
        ctx = build_template_context(
            tpl,
            _merged_job(),
            branding=BrandingConfig(
                company_name="Test Survey Co",
                logo_path=tmp_path / "does_not_exist.png",
                primary_color="#1F4E79",
            ),
            job_name="Job",
            project_name="Project",
            floor_name_for=lambda _fid: "Floor 1",
            photo_paths={1: PHOTO_CLOSE, 2: PHOTO_FAR},
            attachments=_attachments(tmp_path),
        )
    assert ctx["logo"] is None
    assert any("does not resolve" in record.message for record in caplog.records)


def test_unresolved_flags_error(tmp_path: Path) -> None:
    _ensure_photos()
    merged = MergedJob(
        aps=[
            MergedAP(
                ap_name="AP-01",
                survey_data=SurveyAP(name="AP-01"),
                photos=MergedPhotoSlots(),
                status=MergedAPStatus.INCOMPLETE,
            ),
        ],
        flags=[
            Flag(
                ap_name="AP-01",
                type=FlagType.MISSING_PHOTO,
                detail="Missing close photo",
            ),
        ],
    )
    with pytest.raises(UnresolvedFlagsError):
        generate_docx(
            merged,
            template_path=TEMPLATE,
            branding=_branding(),
            job_name="Job",
            project_name="Project",
            floor_name_for=lambda _fid: "—",
            photo_paths={},
            attachments=_attachments(tmp_path),
        )


def test_empty_merged_job_error(tmp_path: Path) -> None:
    with pytest.raises(EmptyMergedJobError):
        generate_docx(
            MergedJob(),
            template_path=TEMPLATE,
            branding=_branding(),
            job_name="Job",
            project_name="Project",
            floor_name_for=lambda _fid: "—",
            photo_paths={},
            attachments=_attachments(tmp_path),
        )


def test_missing_template_error(tmp_path: Path) -> None:
    _ensure_photos()
    with pytest.raises(MissingTemplateError):
        generate_docx(
            _merged_job(),
            template_path=tmp_path / "missing.docx",
            branding=_branding(),
            job_name="Job",
            project_name="Project",
            floor_name_for=lambda _fid: "—",
            photo_paths={1: PHOTO_CLOSE, 2: PHOTO_FAR},
            attachments=_attachments(tmp_path),
        )


def test_no_attachments_error() -> None:
    _ensure_photos()
    with pytest.raises(NoAttachmentsError):
        generate_docx(
            _merged_job(),
            template_path=TEMPLATE,
            branding=_branding(),
            job_name="Job",
            project_name="Project",
            floor_name_for=lambda _fid: "—",
            photo_paths={1: PHOTO_CLOSE, 2: PHOTO_FAR},
            attachments=[],
        )


def test_missing_photo_slots_with_override_still_renders(tmp_path: Path) -> None:
    _ensure_photos()
    merged = MergedJob(
        aps=[
            MergedAP(
                ap_name="AP-01",
                survey_data=SurveyAP(name="AP-01", model="MR46"),
                photos=MergedPhotoSlots(),
                status=MergedAPStatus.INCOMPLETE,
            ),
        ],
        flags=[
            Flag(
                ap_name="AP-01",
                type=FlagType.MISSING_PHOTO,
                detail="No photos",
                override_reason="Not accessible",
            ),
        ],
    )
    docx_bytes = generate_docx(
        merged,
        template_path=TEMPLATE,
        branding=_branding(),
        job_name="Job",
        project_name="Project",
        floor_name_for=lambda _fid: "—",
        photo_paths={},
        attachments=_attachments(tmp_path),
    )
    out = tmp_path / "out.docx"
    out.write_bytes(docx_bytes)
    doc = Document(out)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Not provided" in text
    assert "Not accessible" in text
